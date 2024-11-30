from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

class WordDocument:
  _columns: int
  _font: tuple[str, int]
  _doc: Document

  def __init__(self, columns: int, font: tuple[str, int]):
    self._columns = columns
    self._font = font
    self._setup_document()

  def add_header(self, header: str):
    p = self._doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(header)
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

  def add_content(self, text: str):
    p = self._doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    run = p.add_run(text)
    run.font.size = Pt(self._font[1])
    run.font.name = self._font[0]

  def save(self, path: str):
    self._doc.save(path)

  def _setup_document(self):
    doc = Document()
    section = doc.sections[0]

    # Set page size to A4
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # Set page margins
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)

    # Set columns number to default document section
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(self._columns))

    self._doc = doc